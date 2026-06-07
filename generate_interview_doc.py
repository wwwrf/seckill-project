#!/usr/bin/env python3
"""生成秒杀系统面试问答文档（.docx）。

用法:
  python generate_interview_doc.py
  python generate_interview_doc.py --output 面试问答-秒杀系统.docx
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path


try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "缺少依赖 python-docx，请先执行: pip install python-docx"
    ) from exc


def add_bullets(document: Document, lines: list[str]) -> None:
    for line in lines:
        document.add_paragraph(line, style="List Bullet")


def build_document(output: Path) -> None:
    doc = Document()

    doc.add_heading("电商秒杀系统面试问答（基于当前代码实现）", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        "说明：本答案严格依据当前仓库实现，不按简历措辞臆测。面试时可优先讲“已实现”，再补“改进项”。"
    )

    doc.add_heading("面试前先纠偏（建议主动说明）", level=2)
    add_bullets(
        doc,
        [
            "当前代码中订单超时延迟消息是 30 分钟（RocketMQ delay level=16），不是 15 分钟。",
            "一人一单数据库唯一索引是 (user_id, activity_id)，不是 (user_id, sku_id)。",
            "限流是应用内令牌桶（内存 go-cache + x/time/rate），不是 Redis 限流。",
        ],
    )

    qa = [
        {
            "title": "1) Redis 预扣成功但 MQ 挂了怎么办？重复消费如何处理？延迟消息丢了怎么办？",
            "points": [
                "主链路先写 Outbox（seckill_tx_tasks），再由调度器发 MQ：Redis 预扣成功后，不是直接依赖 MQ 在线，而是先把主消息持久化到 MySQL。",
                "主消息调度器每 1 秒轮询待发送任务，发送失败做指数退避重试，超过阈值进入 dead；dead 再由自动重放器重新入队。",
                "因此 MQ 短暂宕机时不会丢主消息，只会拉长下单落库时延。",
                "重复消费通过数据库幂等兜底：orders 表有 uk_user_activity + uk_order_no。重复消息触发唯一键冲突后 ACK，不会重复建单。",
                "延迟消息发送失败时会写 order_timeout_tasks（延迟消息 Outbox），由补偿器定时重投，失败超阈值进 dead，再由 dead 重放器回捞。",
                "如果出现极端场景（例如长期不可达）日志中已明确“需人工对账”，说明系统语义是高可恢复的 at-least-once，而非绝对零人工。",
            ],
        },
        {
            "title": "2) Lua 里用了哪些 Key？结构是什么？pending 做什么？什么时候清？Redis 挂了会不会库存失控？",
            "points": [
                "预扣 Lua 三个核心 Key：",
                "seckill:stock:{activity}:{product} -> STRING，库存数。",
                "seckill:purchased:{activity}:{product} -> SET，已购 user_id 集合。",
                "seckill:pending:{activity} -> HASH，field=order_no，value={ts,uid,pid} JSON。",
                "单脚本原子执行 EXISTS/SISMEMBER/GET/DECR/SADD/HSET，保证“扣库存+写已购+写 pending”无并发窗口。",
                "pending 作用是记录“已预扣但未建单”的中间态，供异步链路和定时补偿对账。",
                "pending 清理时机：",
                "链路1建单成功后 HDEL；",
                "重复消费命中幂等后也会尝试 HDEL；",
                "Cron 扫描超时 pending（>5 分钟）时，若 DB 无订单则执行回滚 Lua（INCR+SREM+HDEL）；",
                "Cron 垃圾清理也会删除“DB 已有单”或“已作废标记”的残留 pending。",
                "Redis 挂掉时系统是 fail-close：秒杀接口直接报系统忙，不会绕过 Redis 直接建单，因此不会因为降级导致超卖失控；代价是可用性下降。",
                "你这版代码确实存在 Redis 单点可用性风险（local.yaml 是单实例），生产建议上 Sentinel/Cluster + 多 AZ + 熔断降级策略。",
                "为什么不选分段锁/队列模型作为主方案：本实现把一致性核心放在 Redis Lua 原子脚本，RTT 更低、实现更直接；但它依赖 Redis 可用性，这一点要在面试里承认并给 HA 方案。",
            ],
        },
        {
            "title": "3) 唯一索引 + 条件更新细节：索引建在哪？SQL 怎么写？热点行怎么办？",
            "points": [
                "唯一索引：orders.uk_user_activity = (user_id, activity_id)，另有 uk_order_no(order_no)。",
                "不是 (user_id, sku_id)。当前语义是“同一活动一人一单”。",
                "库存条件更新 SQL 语义：UPDATE seckill_activities SET available_stock = available_stock - 1, version = version + 1 WHERE id = ? AND available_stock > 0。",
                "通过 RowsAffected == 0 判断库存不足/并发失败，返回 ErrStockNotEnough。",
                "高并发下 seckill_activities 单活动单行天然是热点行，当前靠 Redis 预扣把绝大多数请求挡在 DB 外。",
                "库存分桶（bucket stock）当前未实现，这是已知可扩展优化点。",
            ],
        },
        {
            "title": "4) 用户级 + IP 级限流怎么做？在哪里做？被限流返回什么？有没有黑名单/滑窗？热门商品单独限流？",
            "points": [
                "实现方式：golang.org/x/time/rate 令牌桶 + go-cache 存每个 user/ip 的 limiter。",
                "不是 Redis 限流，也不是漏桶。",
                "位置：应用层中间件，不是网关层。",
                "链路顺序是 Trace -> JWT -> 用户限流；秒杀下单接口额外叠加 IP 前置限流。",
                "被限流返回 HTTP 429 + 业务码 42900。",
                "当前代码没有设置 Retry-After Header（注释里提到，但实现未写）。",
                "用户体验优化点：令牌桶本身允许 burst，减少误伤；并且用户级限流放在 JWT 后，避免 NAT 下纯 IP 限流误伤。",
                "黑名单机制、滑动窗口、按活动/商品维度的热点限流，当前都未实现。",
            ],
        },
        {
            "title": "5) Outbox + 调度器 + dead 重放：表设计、状态机、轮询/CDC、扫表压力、去重语义",
            "points": [
                "主消息 Outbox 表：seckill_tx_tasks，关键字段含 order_no(唯一)、payload、status、retry_count、next_retry_at、last_error、sent_at。",
                "延迟消息 Outbox 表：order_timeout_tasks，状态与重试字段同类设计。",
                "状态机：pending(0) -> sent(1)；失败重试超过阈值 -> dead(2)；dead 可被 requeue 回 pending。",
                "调度方式是轮询，不是 CDC。主消息调度周期 1s，延迟消息补偿周期 30s。",
                "扫表压力控制手段：按 status + next_retry_at 复合索引过滤、按 batch 限制单次拉取。",
                "不重复发送不是强保证：发送成功但 MarkSent 失败会导致再次发送；多实例调度也可能并发抓到同一 pending。",
                "因此整体语义是 at-least-once，依赖消费端幂等（唯一索引）实现“最终不重单”。",
            ],
        },
        {
            "title": "6) RocketMQ 延迟消息精度、延迟漂移、回滚路径和幂等",
            "points": [
                "本项目使用 RocketMQ 固定延迟级别：delay level 16 = 30 分钟。",
                "精度是“延迟级别 + Broker 调度粒度”语义，不是毫秒级精确触发。",
                "如果 30 分钟延迟漂移到更久，结果是“超时取消触发变晚”，不会直接造成超卖，但会延长订单 pending 时间。",
                "回滚路径是先 DB 后 Redis：先把订单从 pending 改 canceled 并回补 DB 库存，再做 Redis INCR/SREM。",
                "DB 侧幂等通过 WHERE status=pending + RowsAffected 控制；重复延迟消息不会反复取消。",
                "Redis 回滚脚本本身不带业务幂等键，但在 DB 状态门控下通常只执行一次。",
                "若延迟消息“发送失败/丢失”，由 order_timeout_tasks Outbox + 补偿器 + dead 重放兜底。",
            ],
        },
        {
            "title": "7) k6 脚本、arrival-rate 与 VU 区别、为何不用 wrk/JMeter、L1/L2/L3、瓶颈",
            "points": [
                "全链路脚本是 benchmark_e2e_tps.js：下单 -> 轮询 -> 支付，执行器用 ramping-arrival-rate。",
                "arrival-rate 控制“请求到达速率目标”，VU 是“执行并发工人上限”；当 VU 顶满且仍达不到目标时会出现 dropped_iterations。",
                "选 k6 的原因是：多步骤业务脚本表达能力强（含 token、轮询、业务码统计）、支持 arrival-rate 压力模型、易导出标准化 summary。",
                "wrk 更偏单接口高吞吐，复杂业务编排弱；JMeter 可做但配置维护成本更高。",
                "L1/L2/L3 在本项目报告里是三档交易压力级别（例如 stage3=420/700/1100），用来区分稳定区、退化区、拥塞区。",
                "报告显示：读链路可到 10k RPS 仍低错误；交易链路在高压下出现 poll_timeout 和 dropped_iterations，TPS 在拥塞区反降。",
            ],
        },
        {
            "title": "8) Redis 挂 30 秒、MQ 堆积 100 万、真正瓶颈在哪里？",
            "points": [
                "Redis 挂 30 秒：秒杀入口会大量返回系统忙（fail-close），不会放量到 DB 造成超卖；恢复后请求可继续。",
                "这 30 秒内主流程吞吐显著下降，属于“保一致性牺牲可用性”的取舍。",
                "MQ 堆积 100 万时，系统会进入明显异步排队：订单成功可见性变慢，轮询超时增多。",
                "现有可用手段：提高消费者并发与实例数、分离 seckill/poll/pay 限流、通过 outbox/retry/dead 回放保障消息最终可达。",
                "真正瓶颈通常不在读缓存层，而在“异步建单消费能力 + DB 写热点 + 轮询放大效应”的组合。",
                "从现有压测结果看，交易链路先退化，读链路后退化，这与架构分层结论一致。",
            ],
        },
    ]

    for item in qa:
        doc.add_heading(item["title"], level=2)
        add_bullets(doc, item["points"])

    doc.add_heading("附录：关键实现位置（便于面试时快速定位）", level=2)
    add_bullets(
        doc,
        [
            "秒杀主流程：internal/service/seckill_service.go",
            "Redis Lua 与 pending/回滚：internal/repository/cache/seckill_cache.go",
            "Redis Key 约定：internal/model/e/keys.go",
            "主消息 Outbox：internal/model/seckill_tx_task.go + internal/repository/seckill_tx_task_repo.go",
            "主消息调度与 dead 重放：internal/cron/tx_message_dispatcher.go + internal/cron/tx_task_replayer.go",
            "MQ 消费链路（建单/超时取消）：internal/mq/consumer.go",
            "延迟消息 Outbox 补偿：internal/model/order_timeout_task.go + internal/cron/timeout_message_compensator.go",
            "订单幂等与库存扣减：internal/repository/order_repo.go",
            "限流中间件：internal/middleware/middleware.go",
            "压测脚本：loadtest/benchmark_e2e_tps.js 与 loadtest/LOAD_TEST_REPORT_v2_20260331.md",
        ],
    )

    doc.add_heading("面试话术建议（加分）", level=2)
    add_bullets(
        doc,
        [
            "先讲已落地机制，再主动讲边界和改进项，可信度更高。",
            "明确语义是 at-least-once + 业务幂等，不要说 exactly-once。",
            "主动承认 Redis 单点风险，并给出 Sentinel/Cluster、多机房、熔断降级方案。",
            "对压测结论区分“稳定容量”和“极限峰值”，避免只报峰值。",
        ],
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成秒杀系统面试问答 docx")
    parser.add_argument(
        "--output",
        default="面试问答-秒杀系统.docx",
        help="输出 docx 文件路径，默认：面试问答-秒杀系统.docx",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output = Path(args.output)
    build_document(output)
    print(f"文档已生成: {output.resolve()}")


if __name__ == "__main__":
    main()
